"""
Generate .docx report for BTL ATBMTT - FIT4012
Secure System Upgrade Challenge
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ====================================================================
# GLOBAL STYLES
# ====================================================================

# --- Margins: Left 3cm, Right/Top/Bottom 2cm ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# --- Default font ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.space_before = Pt(6)
style.paragraph_format.space_after = Pt(6)

# Set East Asian font fallback
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def add_heading_styled(text, level=1):
    """Add a heading with Times New Roman, uppercase, bold, size 15-16."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(15)
        else:
            run.font.size = Pt(14)
    return h


def add_para(text, bold=False, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent_left=None, space_before=6, space_after=6):
    """Add a paragraph with standard formatting."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if indent_left is not None:
        p.paragraph_format.left_indent = Cm(indent_left)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p


def add_para_multi(segments, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   indent_left=None, space_before=6, space_after=6):
    """Add paragraph with mixed bold/normal segments.
    segments: list of (text, bold) tuples
    """
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if indent_left is not None:
        p.paragraph_format.left_indent = Cm(indent_left)
    for text, bold in segments:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.bold = bold
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p


def add_code_block(code, indent_left=1):
    """Add a code block as monospaced text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    if indent_left is not None:
        p.paragraph_format.left_indent = Cm(indent_left)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return p


def add_bullet_indent(text, indent=1.0, size=13):
    """Add indented text (replaces bullet points per user request)."""
    return add_para(text, size=size, indent_left=indent)


def add_table_with_data(headers, rows, col_widths=None, header_color="2F5496"):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(header)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(cell_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table


# ====================================================================
# TRANG BÌA
# ====================================================================

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC CÔNG NGHIỆP HÀ NỘI")
run.font.name = 'Times New Roman'
run.font.size = Pt(13)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("--------***--------")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BÁO CÁO BÀI TẬP LỚN\nFIT4012 - AN TOÀN VÀ BẢO MẬT THÔNG TIN")
run.font.name = 'Times New Roman'
run.font.size = Pt(18)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ĐỀ TÀI: SECURE SYSTEM UPGRADE CHALLENGE\nMODERN PASSWORD AUTHENTICATION SYSTEM")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Giảng viên hướng dẫn: TS. Nguyễn Văn A\nNhóm sinh viên thực hiện: Nhóm 1\nMã học phần: FIT4012\nNăm học: 2025-2026")
run.font.name = 'Times New Roman'
run.font.size = Pt(13)

doc.add_page_break()

# ====================================================================
# MỤC LỤC (Manual Table of Contents)
# ====================================================================

add_heading_styled("MỤC LỤC", level=1)

toc_items = [
    ("1.", "Giới thiệu bài toán", 3),
    ("2.", "Mục tiêu bảo mật", 4),
    ("3.", "Threat model", 5),
    ("4.", "Kiến trúc hệ thống", 6),
    ("5.", "Thiết kế giao thức và luồng xử lý", 7),
    ("6.", "Thuật toán và thư viện sử dụng", 8),
    ("7.", "Mô tả chức năng đã cài đặt", 9),
    ("8.", "Phân tích mã nguồn", 10),
    ("9.", "Kiểm thử chức năng", 11),
    ("10.", "Kiểm thử bảo mật", 12),
    ("11.", "Benchmark hiệu năng", 13),
    ("12.", "Bảng kế thừa và nâng cấp", 14),
    ("13.", "Kết luận và hướng phát triển", 15),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(f"{num} {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = False

doc.add_page_break()

# ====================================================================
# CHƯƠNG 1: GIỚI THIỆU BÀI TOÁN
# ====================================================================

add_heading_styled("CHƯƠNG 1: GIỚI THIỆU BÀI TOÁN", level=1)

add_heading_styled("1.1. Bối cảnh thực tế", level=2)

add_para(
    "Trong thời đại số hóa hiện nay, các hệ thống xác thực người dùng đóng vai trò then chốt "
    "trong việc bảo vệ thông tin cá nhân và tài nguyên số. Các cuộc tấn công nhắm vào cơ chế "
    "xác thực ngày càng tinh vi, từ tấn công brute-force, dictionary attack cho đến khai thác "
    "lỗ hổng lưu trữ mật khẩu dạng plaintext. Nhiều vụ rò rỉ dữ liệu lớn trên thế giới bắt nguồn "
    "từ việc hệ thống lưu trữ mật khẩu không an toàn, cho phép kẻ tấn công truy cập trái phép "
    "vào hàng triệu tài khoản người dùng."
)
add_para(
    "Bài toán Secure System Upgrade Challenge được đặt ra trong bối cảnh một hệ thống xác thực "
    "cũ sử dụng các thuật toán mã hóa lỗi thời như SHA-256 (băm nhanh, không có salt) và TripleDES "
    "(mã hóa hai chiều, có thể giải mã ngược). Yêu cầu đặt ra là nâng cấp toàn diện hệ thống này "
    "lên một giải pháp xác thực hiện đại, an toàn, sử dụng thuật toán băm chuyên dụng Argon2id "
    "kết hợp với cơ chế khóa tài khoản và ghi log chi tiết."
)

add_heading_styled("1.2. Mục tiêu của hệ thống", level=2)

add_para(
    "Hệ thống được xây dựng nhằm đạt được các mục tiêu sau:"
)
add_bullet_indent("- Cung cấp API xác thực an toàn gồm ba chức năng chính: đăng ký (Register), đăng nhập (Login) và đổi mật khẩu (Change Password).", indent=1.0)
add_bullet_indent("- Sử dụng thuật toán băm Argon2id để lưu trữ mật khẩu an toàn, đảm bảo mỗi mật khẩu được tạo salt ngẫu nhiên riêng biệt.", indent=1.0)
add_bullet_indent("- Triển khai cơ chế khóa tài khoản (Account Lockout) và backoff lũy thừa để chống tấn công brute-force.", indent=1.0)
add_bullet_indent("- Ghi log chi tiết tất cả sự kiện xác thực phục vụ kiểm toán và truy vết.", indent=1.0)
add_bullet_indent("- Xây dựng bộ kiểm thử và benchmark để đánh giá hiệu năng và độ an toàn của hệ thống.", indent=1.0)

add_heading_styled("1.3. Lý do cần bảo mật", level=2)

add_para(
    "Việc bảo mật hệ thống xác thực là vô cùng quan trọng vì những lý do sau:"
)
add_bullet_indent("- Bảo vệ danh tính người dùng: Mật khẩu là lớp bảo vệ đầu tiên cho tài khoản người dùng. Nếu mật khẩu bị lộ, kẻ tấn công có thể chiếm đoạt toàn bộ tài khoản.", indent=1.0)
add_bullet_indent("- Ngăn chặn leo thang đặc quyền: Thông qua tài khoản bị chiếm đoạt, kẻ tấn công có thể leo thang đặc quyền để truy cập các tài nguyên nhạy cảm hơn.", indent=1.0)
add_bullet_indent("- Đáp ứng yêu cầu pháp lý: Nhiều quy định bảo vệ dữ liệu (GDPR, Nghị định 13/2023/NĐ-CP) yêu cầu lưu trữ mật khẩu an toàn.", indent=1.0)
add_bullet_indent("- Bảo vệ uy tín tổ chức: Một vụ rò rỉ dữ liệu có thể gây thiệt hại lớn về uy tín và tài chính cho tổ chức.", indent=1.0)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 2: MỤC TIÊU BẢO MẬT
# ====================================================================

add_heading_styled("CHƯƠNG 2: MỤC TIÊU BẢO MẬT", level=1)

add_heading_styled("2.1. Tính bí mật (Confidentiality)", level=2)

add_para(
    "Mật khẩu người dùng không được lưu trữ dưới dạng plaintext trong cơ sở dữ liệu. "
    "Thay vào đó, hệ thống sử dụng thuật toán băm Argon2id để biến đổi mật khẩu thành chuỗi hash, "
    "đảm bảo ngay cả khi cơ sở dữ liệu bị xâm phạm, kẻ tấn công cũng không thể khôi phục được "
    "mật khẩu gốc. Mỗi mật khẩu được tạo salt ngẫu nhiên 16 byte riêng biệt, nhúng trực tiếp vào "
    "chuỗi hash, ngăn chặn tấn công Rainbow Table."
)

add_heading_styled("2.2. Tính toàn vẹn (Integrity)", level=2)

add_para(
    "Hệ thống đảm bảo tính toàn vẹn của dữ liệu xác thực thông qua việc sử dụng các thuật toán "
    "băm có tính chất một chiều (one-way hash). Mọi thay đổi trái phép đối với dữ liệu người dùng "
    "trong cơ sở dữ liệu sẽ dẫn đến việc xác thực thất bại. Ngoài ra, log sự kiện được ghi lại "
    "để phát hiện các hành vi bất thường."
)

add_heading_styled("2.3. Tính xác thực (Authentication)", level=2)

add_para(
    "Hệ thống xác thực người dùng thông qua việc kiểm tra mật khẩu bằng thuật toán Argon2id. "
    "Chỉ người dùng có mật khẩu đúng mới có thể đăng nhập thành công. "
    "Ngoài ra, cơ chế khóa tài khoản tạm thời được kích hoạt khi phát hiện nhiều lần đăng nhập sai, "
    "ngăn chặn tấn công brute-force và dictionary attack."
)

add_heading_styled("2.4. Tính sẵn sàng (Availability)", level=2)

add_para(
    "Hệ thống được thiết kế với kiến trúc API nhẹ dựa trên FastAPI và SQLite, đảm bảo khả năng "
    "phục vụ liên tục. Cơ chế backoff lũy thừa giúp giảm tải cho server khi có nhiều yêu cầu sai "
    "liên tiếp, bảo vệ tài nguyên hệ thống. Thời gian khóa tài khoản có giới hạn (15 phút) để "
    "người dùng hợp pháp có thể thử lại sau đó."
)

add_heading_styled("2.5. Khả năng truy vết (Accountability)", level=2)

add_para(
    "Mọi sự kiện xác thực đều được ghi lại trong file auth.log với định dạng chuẩn bao gồm: "
    "thời gian, tên người dùng, hành động, lý do và địa chỉ IP. Log được thiết kế không chứa "
    "mật khẩu, password hash hay secret key, đảm bảo an toàn khi lưu trữ và kiểm toán."
)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 3: THREAT MODEL
# ====================================================================

add_heading_styled("CHƯƠNG 3: THREAT MODEL", level=1)

add_heading_styled("3.1. Tài sản cần bảo vệ", level=2)

add_para("Các tài sản cần được bảo vệ trong hệ thống bao gồm:")
add_bullet_indent("- Thông tin tài khoản người dùng: username, password_hash, trạng thái khóa.", indent=1.0)
add_bullet_indent("- Mật khẩu gốc (plaintext) của người dùng.", indent=1.0)
add_bullet_indent("- Log xác thực chứa thông tin lịch sử đăng nhập.", indent=1.0)
add_bullet_indent("- Cơ sở dữ liệu (app.db) chứa toàn bộ dữ liệu người dùng.", indent=1.0)
add_bullet_indent("- File log (auth.log) chứa thông tin sự kiện.", indent=1.0)

add_heading_styled("3.2. Người dùng hợp lệ", level=2)

add_para("Các tác nhân được phép truy cập hệ thống:")
add_bullet_indent("- Người dùng cuối (End User): Đăng ký tài khoản, đăng nhập, đổi mật khẩu.", indent=1.0)
add_bullet_indent("- Hệ thống backend (Server): Xử lý request, tương tác với database.", indent=1.0)
add_bullet_indent("- Quản trị viên (Administrator): Kiểm tra log, quản lý tài khoản.", indent=1.0)

add_heading_styled("3.3. Tác nhân tấn công giả định", level=2)

add_para("Các tác nhân tấn công mà hệ thống giả định phải đối mặt:")
add_bullet_indent("- Kẻ tấn công brute-force: Thử hàng loạt mật khẩu để đăng nhập vào tài khoản.", indent=1.0)
add_bullet_indent("- Kẻ tấn công có quyền truy cập database: Có thể đọc dữ liệu bảng users.", indent=1.0)
add_bullet_indent("- Kẻ tấn công dictionary: Sử dụng danh sách mật khẩu phổ biến để dò tìm.", indent=1.0)
add_bullet_indent("- Kẻ tấn công MITM (Man-in-the-Middle): Nghe lén lưu lượng mạng (nếu không dùng HTTPS).", indent=1.0)
add_bullet_indent("- Kẻ tấn công nội bộ (Insider): Có thể truy cập log hoặc database.", indent=1.0)

add_heading_styled("3.4. Các nguy cơ bảo mật", level=2)

add_para("Bảng dưới đây liệt kê các nguy cơ bảo mật chính:")

risk_rows = [
    ["R01", "Lộ mật khẩu plaintext", "Database bị xâm phạm", "Nghiêm trọng"],
    ["R02", "Brute-force đăng nhập", "Tấn công liên tục vào /login", "Cao"],
    ["R03", "Đọc log chứa thông tin nhạy cảm", "File log bị lộ", "Trung bình"],
    ["R04", "Tấn công Replay", "Request bị bắt và phát lại", "Trung bình"],
    ["R05", "Khai thác tham số đầu vào", "Injection qua username/password", "Cao"],
    ["R06", "Tấn công từ chối dịch vụ (DoS)", "Gửi quá nhiều request", "Thấp"],
]
add_table_with_data(
    ["Mã", "Nguy cơ", "Mô tả", "Mức độ"],
    risk_rows,
    col_widths=[1.5, 4, 6, 2.5]
)

add_heading_styled("3.5. Các giả định và giới hạn của hệ thống", level=2)

add_para("Các giả định khi thiết kế hệ thống:")
add_bullet_indent("- Hệ thống hoạt động trong môi trường mạng nội bộ (Intranet) hoặc có HTTPS.", indent=1.0)
add_bullet_indent("- Người dùng được khuyến khích sử dụng mật khẩu mạnh (tối thiểu 6 ký tự).", indent=1.0)
add_bullet_indent("- Hệ điều hành và hạ tầng mạng được bảo vệ cơ bản.", indent=1.0)

add_para("Các giới hạn:")
add_bullet_indent("- Hệ thống chưa triển khai xác thực hai yếu tố (2FA).", indent=1.0)
add_bullet_indent("- Chưa hỗ trợ khôi phục mật khẩu qua email.", indent=1.0)
add_bullet_indent("- Chưa sử dụng HTTPS trong môi trường phát triển (cần cấu hình thêm).", indent=1.0)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 4: KIẾN TRÚC HỆ THỐNG
# ====================================================================

add_heading_styled("CHƯƠNG 4: KIẾN TRÚC HỆ THỐNG", level=1)

add_heading_styled("4.1. Thành phần hệ thống", level=2)

add_para(
    "Hệ thống bao gồm hai thành phần chính, hoạt động độc lập nhưng tích hợp với nhau:"
)

add_para_multi([
    ("Thành phần 1 - Backend xác thực (Python FastAPI): ", True),
    ("Đây là thành phần cốt lõi của hệ thống, cung cấp các API xác thực người dùng. "
     "Bao gồm các module: routes/auth.py (xử lý request), security/password_hasher.py (băm mật khẩu), "
     "security/lockout.py (chính sách khóa tài khoản), logger.py (ghi log), models.py (định nghĩa dữ liệu), "
     "db.py (kết nối CSDL). Hệ thống sử dụng SQLite làm cơ sở dữ liệu.", False),
])

add_para_multi([
    ("Thành phần 2 - Công cụ kiểm toán an ninh (C# .NET): ", True),
    ("Một ứng dụng console riêng biệt (Program.cs + SecurityBenchmark) dùng để chạy benchmark "
     "thuật toán và kiểm thử tích hợp các tính năng bảo mật như brute-force mitigation, "
     "change password lifecycle. Sử dụng BCrypt, Argon2 và JWT cho các bài kiểm tra.", False),
])

add_heading_styled("4.2. Sơ đồ kiến trúc tổng quan", level=2)

add_para(
    "Vui lòng dán mã PlantUML sau vào https://editor.plantuml.com/ để xem sơ đồ kiến trúc:"
)

architecture_plantuml = """@startuml
!theme plain
skinparam backgroundColor #FFFFFF
skinparam componentStyle rectangle

title Kiến trúc tổng quan hệ thống

package "Client" {
  [HTTP Client\n(Python requests)] as Client
}

package "Backend - FastAPI Application" {
  [Router\n(app/router.py)] as Router
  [Auth Routes\n(app/routes/auth.py)] as AuthRoutes
  [Password Hasher\napp/security/password_hasher.py] as Hasher
  [Lockout Policy\napp/security/lockout.py] as Lockout
  [Logger\napp/logger.py] as Logger
  [Models\napp/models.py] as Models
  [Database\napp/db.py] as DB
}

package "Storage" {
  [SQLite DB\napp.db] as SQLite
  [Auth Log\nauth.log] as AuthLog
}

package "Security Audit Tool\n(C# .NET)" {
  [Program.cs\nTC01_BruteForceTests\nTC02_ChangePasswordTests\nTC03_SecurityAuditTests] as Audit
  [BenchmarkProgram\nBCrypt / Argon2] as Benchmark
}

Client --> Router : HTTP POST\n/register, /login\n/change-password
Router --> AuthRoutes
AuthRoutes --> Hasher : hash_password()\nverify_password()
AuthRoutes --> Lockout : compute_backoff()\ncompute_locked_until()
AuthRoutes --> Logger : log_action()
AuthRoutes --> Models : User, LoginLog
Models --> DB
DB --> SQLite
Logger --> AuthLog
Audit .up.> Client : simulates\nHTTP requests
Benchmark --> Audit

@enduml"""

add_code_block(architecture_plantuml, indent_left=0)

add_para(
    "Hình 1: Kiến trúc tổng quan hệ thống - Minh họa mối quan hệ giữa Client, Backend FastAPI, "
    "cơ sở dữ liệu và công cụ kiểm toán C#."
)

add_heading_styled("4.3. Mô tả vai trò từng thành phần", level=2)

add_para_multi([
    ("FastAPI Application (app/main.py): ", True),
    ("Điểm khởi đầu của ứng dụng, cấu hình router và tự động tạo bảng CSDL khi khởi động.", False),
])

add_para_multi([
    ("Router (app/router.py): ", True),
    ("Tập hợp tất cả các API endpoint, hiện tại bao gồm router xác thực.", False),
])

add_para_multi([
    ("Auth Routes (app/routes/auth.py): ", True),
    ("Xử lý logic nghiệp vụ cho các endpoint /register, /login, /change-password. "
     "Bao gồm kiểm tra đầu vào, xác thực mật khẩu, quản lý khóa tài khoản và ghi log.", False),
])

add_para_multi([
    ("Password Hasher (app/security/password_hasher.py): ", True),
    ("Module chịu trách nhiệm băm và xác minh mật khẩu sử dụng Argon2id với salt 16 byte, "
     "time_cost=2, memory_cost=102400KB (~100MB), parallelism=8.", False),
])

add_para_multi([
    ("Lockout Policy (app/security/lockout.py): ", True),
    ("Định nghĩa chính sách khóa tài khoản: tối đa 5 lần sai, thời gian khóa 15 phút, "
     "backoff lũy thừa (1s, 2s, 4s, 8s, ... tối đa 30s).", False),
])

add_para_multi([
    ("Logger (app/logger.py): ", True),
    ("Cấu hình logging ghi sự kiện vào file auth.log với định dạng [timestamp] User: ... Action: ...", False),
])

add_para_multi([
    ("Models (app/models.py): ", True),
    ("Định nghĩa ORM User (id, username, password_hash, failed_attempts, locked_until, created_at) "
     "và LoginLog (id, user_id, timestamp, status, ip_address).", False),
])

add_para_multi([
    ("Security Audit Tool (C# .NET): ", True),
    ("Chương trình console thực thi 3 bộ kiểm thử: TC01 - BruteForceTests, "
     "TC02 - ChangePasswordTests, TC03 - SecurityAuditTests, kèm benchmark thuật toán.", False),
])

doc.add_page_break()

# ====================================================================
# CHƯƠNG 5: THIẾT KẾ GIAO THỨC VÀ LUỒNG XỬ LÝ
# ====================================================================

add_heading_styled("CHƯƠNG 5: THIẾT KẾ GIAO THỨC VÀ LUỒNG XỬ LÝ", level=1)

add_heading_styled("5.1. Sơ đồ Sequence Diagram", level=2)

add_para(
    "Vui lòng dán mã PlantUML sau vào https://editor.plantuml.com/ để xem sơ đồ tuần tự:"
)

sequence_plantuml = """@startuml
!theme plain
skinparam backgroundColor #FFFFFF
actor User
participant "Client" as C
participant "FastAPI\n/login" as API
participant "Lockout\nPolicy" as L
participant "Password\nHasher" as H
participant "Database" as DB
participant "Logger" as Log

== Đăng ký (Register) ==
User -> C: username, password
C -> API: POST /register
API -> API: validate input\n(username not empty,\npassword >= 6 chars)
API -> DB: check existing username
DB --> API: not found
API -> H: hash_password(password)
H --> API: $argon2id$...
API -> DB: INSERT new user
API -> Log: log_action("Register Success")
API --> C: 200 {id, username}
C --> User: Success

== Đăng nhập thành công ==
User -> C: username, password
C -> API: POST /login
API -> DB: SELECT user by username
DB --> API: user data
API -> L: check locked_until > now?
L --> API: not locked
API -> H: verify_password(hash, password)
H --> API: True (match)
API -> DB: reset failed_attempts=0
API -> Log: log_action("Login Success")
API --> C: {status: "success"}
C --> User: Success

== Đăng nhập sai -> Lockout ==
User -> C: username, wrong password (x5)
C -> API: POST /login (5 times)
API -> H: verify_password(hash, wrong)
H --> API: False
API -> DB: failed_attempts +=1
alt failed_attempts < 5
  API -> L: compute_backoff_seconds()
  API -> DB: locked_until = now + backoff
  API -> Log: log_action("Login Failed")
else failed_attempts >= 5
  API -> L: compute_locked_until()
  API -> DB: locked_until = now + 15min
  API -> Log: log_action("Account Locked")
end
API --> C: {status: "fail"}
C --> User: Invalid Credentials

@enduml"""

add_code_block(sequence_plantuml, indent_left=0)

add_para(
    "Hình 2: Sequence Diagram - Luồng xử lý đăng ký, đăng nhập thành công và cơ chế lockout."
)

add_heading_styled("5.2. Mô tả các bước xử lý", level=2)

add_para_multi([
    ("Bước 1 - Khởi tạo: ", True),
    ("Người dùng gửi request HTTP đến server. FastAPI nhận request, phân tích JSON body và chuyển đến router tương ứng.", False),
])

add_para_multi([
    ("Bước 2 - Xác thực đầu vào: ", True),
    ("Hệ thống kiểm tra tính hợp lệ của dữ liệu: username không rỗng, password tối thiểu 6 ký tự. "
     "Nếu không hợp lệ, trả về HTTP 400.", False),
])

add_para_multi([
    ("Bước 3 - Kiểm tra tồn tại: ", True),
    ("Đối với đăng ký: kiểm tra username đã tồn tại chưa (HTTP 409 nếu có). "
     "Đối với đăng nhập: nếu không tìm thấy user, trả về 'Invalid credentials' mà không tiết lộ thông tin.", False),
])

add_para_multi([
    ("Bước 4 - Kiểm tra khóa tài khoản: ", True),
    ("Trước khi xác thực mật khẩu, hệ thống kiểm tra locked_until. "
     "Nếu tài khoản đang bị khóa, từ chối ngay lập tức.", False),
])

add_para_multi([
    ("Bước 5 - Xác thực mật khẩu: ", True),
    ("Sử dụng argon2-cffi để xác minh mật khẩu với chuỗi hash. Argon2id tự động giải mã salt từ chuỗi hash.", False),
])

add_para_multi([
    ("Bước 6 - Xử lý kết quả: ", True),
    ("Thành công: reset failed_attempts, xóa locked_until, ghi log Login Success. "
     "Thất bại: tăng failed_attempts, áp dụng backoff/lockout, ghi log Login Failed.", False),
])

add_para_multi([
    ("Bước 7 - Ghi log: ", True),
    ("Mọi sự kiện đều được ghi vào file auth.log với đầy đủ thông tin: timestamp, username, action, reason, IP.", False),
])

doc.add_page_break()

# ====================================================================
# CHƯƠNG 6: THUẬT TOÁN VÀ THƯ VIỆN SỬ DỤNG
# ====================================================================

add_heading_styled("CHƯƠNG 6: THUẬT TOÁN VÀ THƯ VIỆN SỬ DỤNG", level=1)

add_heading_styled("6.1. Thuật toán mã hóa và băm", level=2)

add_para_multi([
    ("Argon2id: ", True),
    ("Đây là thuật toán băm mật khẩu chính của hệ thống. Argon2id là phiên bản kết hợp giữa "
     "Argon2i (chống side-channel) và Argon2d (chống GPU cracking), được công nhận là chuẩn "
     "băm mật khẩu tối ưu nhất hiện nay (chiến thắng cuộc thi Password Hashing Competition 2015). "
     "Hỗ trợ ba tham số cấu hình: time_cost (số lần lặp), memory_cost (bộ nhớ sử dụng), "
     "parallelism (số luồng song song). Hệ thống sử dụng salt ngẫu nhiên 16 byte.", False),
])

add_para_multi([
    ("SHA-256 (Legacy - dùng để so sánh): ", True),
    ("Thuật toán băm nhanh thuộc họ SHA-2, tạo ra hash 256-bit. "
     "Tuy nhiên, tốc độ nhanh là điểm yếu khi dùng cho mật khẩu vì dễ bị brute-force. "
     "Được dùng trong legacy_comparison.js để minh họa sự khác biệt.", False),
])

add_para_multi([
    ("TripleDES (Legacy - dùng để so sánh): ", True),
    ("Thuật toán mã hóa đối xứng 64-bit, mã hóa/giải mã hai chiều. "
     "Hoàn toàn không phù hợp để lưu trữ mật khẩu vì nếu lộ khóa có thể giải mã ngược ra plaintext. "
     "Được dùng trong legacy_comparison.js để minh họa.", False),
])

add_para_multi([
    ("BCrypt (Dùng trong C# Benchmark): ", True),
    ("Thuật toán băm mật khẩu phổ biến dựa trên Blowfish. Được sử dụng trong công cụ kiểm toán C# "
     "để so sánh hiệu năng với Argon2.", False),
])

add_heading_styled("6.2. Cơ chế quản lý khóa", level=2)

add_para(
    "Argon2id tự động tạo salt ngẫu nhiên 16 byte cho mỗi lần băm. Salt được nhúng trực tiếp "
    "vào chuỗi hash dưới định dạng: $argon2id$v=19$m=102400,t=2,p=8$<salt>$<hash>. "
    "Khi xác minh, Argon2id tự động giải mã salt từ chuỗi hash, do đó không cần lưu salt riêng biệt. "
    "Điều này đảm bảo hai người dùng có cùng mật khẩu sẽ có hash hoàn toàn khác nhau."
)

add_heading_styled("6.3. Thư viện lập trình sử dụng", level=2)

lib_rows = [
    ["fastapi==0.115.0", "Python", "Framework web, xây dựng REST API"],
    ["uvicorn[standard]==0.30.6", "Python", "ASGI server, chạy ứng dụng FastAPI"],
    ["SQLAlchemy==2.0.35", "Python", "ORM, tương tác với SQLite database"],
    ["argon2-cffi==23.1.0", "Python", "Thư viện Argon2id password hashing"],
    ["pydantic==2.8.2", "Python", "Xác thực dữ liệu request/response"],
    ["requests==2.32.3", "Python", "HTTP client cho demo/test tự động"],
    ["BCrypt.Net-Next 4.2.0", "C# (.NET)", "BCrypt hashing trong công cụ kiểm toán"],
    ["Konscious.Argon2 1.3.1", "C# (.NET)", "Argon2 hashing trong công cụ kiểm toán"],
    ["System.IdentityModel.Tokens.Jwt", "C# (.NET)", "Tạo/xác thực JWT token"],
]
add_table_with_data(
    ["Thư viện", "Nền tảng", "Mục đích sử dụng"],
    lib_rows,
    col_widths=[4.5, 2.5, 7]
)

add_heading_styled("6.4. Lý do lựa chọn", level=2)

add_para("Bảng so sánh các thuật toán hashing:")

comparison_rows = [
    ["SHA-256", "Băm nhanh (~1µs)", "Không có salt", "Rất dễ (GPU)", "Không phù hợp"],
    ["TripleDES", "Mã hóa 2 chiều", "Khóa bí mật", "Giải mã nếu lộ khóa", "Không phù hợp"],
    ["BCrypt", "Băm chậm (~100ms)", "Salt nhúng", "Khó (cost cao)", "Chấp nhận được"],
    ["Argon2id", "Băm chậm, tùy chỉnh", "Salt 16B nhúng", "Rất khó (memory-hard)", "Tối ưu nhất"],
]
add_table_with_data(
    ["Thuật toán", "Đặc điểm", "Cơ chế salt", "Khả năng chống tấn công", "Đánh giá"],
    comparison_rows,
    col_widths=[2.5, 3.5, 3, 4, 3]
)

add_para(
    "Argon2id được lựa chọn vì: (1) Là thuật toán chiến thắng Password Hashing Competition 2015; "
    "(2) Tính chất memory-hard chống GPU cracking hiệu quả; "
    "(3) Tự động tạo salt ngẫu nhiên, không cần quản lý salt riêng; "
    "(4) Có thể điều chỉnh tham số để cân bằng giữa bảo mật và hiệu năng."
)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 7: MÔ TẢ CHỨC NĂNG ĐÃ CÀI ĐẶT
# ====================================================================

add_heading_styled("CHƯƠNG 7: MÔ TẢ CHỨC NĂNG ĐÃ CÀI ĐẶT", level=1)

add_heading_styled("7.1. Chức năng nền (Base)", level=2)

add_para_multi([
    ("Đăng ký tài khoản (POST /register): ", True),
    ("Cho phép người dùng tạo tài khoản mới với username và password. "
     "Mật khẩu được băm bằng Argon2id trước khi lưu vào database. "
     "Kiểm tra username không rỗng, password tối thiểu 6 ký tự, username không trùng lặp.", False),
])

add_para_multi([
    ("Đăng nhập (POST /login): ", True),
    ("Xác thực người dùng với username và password. "
     "Kiểm tra trạng thái khóa tài khoản trước khi xác minh mật khẩu. "
     "Sử dụng Argon2id để verify password_hash. Reset failed_attempts nếu thành công.", False),
])

add_para_multi([
    ("Đổi mật khẩu (POST /change-password): ", True),
    ("Cho phép người dùng đổi mật khẩu sau khi xác minh mật khẩu cũ. "
     "Mật khẩu mới được tạo hash mới (salt mới). "
     "Yêu cầu mật khẩu mới tối thiểu 6 ký tự và khác mật khẩu cũ.", False),
])

add_heading_styled("7.2. Chức năng nâng cấp (Upgrade)", level=2)

add_para_multi([
    ("Cơ chế khóa tài khoản (Account Lockout): ", True),
    ("Nâng cấp từ hệ thống không có khóa tài khoản. "
     "Sau 5 lần đăng nhập sai, tài khoản bị khóa 15 phút. "
     "Áp dụng backoff lũy thừa (1s, 2s, 4s...) trước khi đạt ngưỡng khóa.", False),
])

add_para_multi([
    ("Log sự kiện xác thực: ", True),
    ("Ghi log chi tiết tất cả sự kiện: đăng ký, đăng nhập thành công, thất bại, khóa tài khoản, đổi mật khẩu. "
     "Log không chứa mật khẩu hay password hash. Mỗi dòng log đều có timestamp và địa chỉ IP.", False),
])

add_para_multi([
    ("Hash mật khẩu bằng Argon2id: ", True),
    ("Nâng cấp từ SHA-256/TripleDES lên Argon2id. "
     "Salt ngẫu nhiên 16 byte cho mỗi hash. Memory cost 100MB chống GPU cracking.", False),
])

add_para_multi([
    ("Không tiết lộ thông tin: ", True),
    ("Khi đăng nhập sai, hệ thống trả về 'Invalid credentials' chung, "
     "không phân biệt 'username không tồn tại' và 'sai mật khẩu' để tránh rò rỉ thông tin.", False),
])

add_heading_styled("7.3. Giao diện (CLI)", level=2)

add_para(
    "Hệ thống hiện tại cung cấp API REST, không có giao diện đồ họa. "
    "Có thể tương tác qua HTTP Client (curl, Postman, requests Python). "
    "Script demo.py cung cấp CLI tự động để kiểm thử các chức năng."
)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 8: PHÂN TÍCH MÃ NGUỒN
# ====================================================================

add_heading_styled("CHƯƠNG 8: PHÂN TÍCH MÃ NGUỒN", level=1)

add_heading_styled("8.1. Cấu trúc thư mục", level=2)

add_para("Cấu trúc thư mục của dự án như sau:")

folder_structure = """D:\BTTL_Son/
├── app/                          # Backend Python FastAPI
│   ├── main.py                   # Khởi tạo ứng dụng FastAPI, tạo bảng DB
│   ├── router.py                 # Tổng hợp các router API
│   ├── db.py                     # Kết nối SQLite database (SQLAlchemy)
│   ├── models.py                 # ORM models: User, LoginLog
│   ├── schemas.py                # Pydantic schemas cho request/response
│   ├── logger.py                 # Cấu hình ghi log ra file auth.log
│   ├── routes/
│   │   └── auth.py               # Logic xử lý register/login/change-password
│   └── security/
│       ├── password_hasher.py    # Argon2id hash & verify
│       └── lockout.py            # Chính sách account lockout & backoff
├── scripts/
│   ├── demo.py                   # Script demo tự động gọi API
│   └── legacy_comparison.js      # So sánh SHA-256, TripleDES, Argon2id
├── Output/                       # Thư mục chứa kết quả đầu ra
├── SecurityBenchmark/            # Module C# benchmark thuật toán
├── Services/                     # Services C# cho kiểm thử
├── Models/                       # Models C#
├── Program.cs                    # Entry point C# (console kiểm toán)
├── SecurityBenchmark.csproj      # Project file .NET 8.0
├── schema.sql                    # SQL schema tham khảo
├── auth.log                      # File log sự kiện xác thực
├── app.db                        # SQLite database
├── requirements.txt              # Python dependencies
├── test_cases.md                 # Danh sách test case
├── test_report.md                # Báo cáo kiểm thử
└── README.md                     # Hướng dẫn chạy ứng dụng"""

add_code_block(folder_structure, indent_left=0)

add_heading_styled("8.2. Các file chính", level=2)

add_para_multi([
    ("app/main.py: ", True),
    ("Tạo ứng dụng FastAPI, include router, tạo bảng CSDL khi startup. "
     "Là điểm vào khi chạy lệnh uvicorn app.main:app.", False),
])

add_para_multi([
    ("app/routes/auth.py: ", True),
    ("Chứa 3 endpoint chính. Sử dụng dependency injection (Depends(get_db)) để quản lý session DB. "
     "Đây là file quan trọng nhất vì chứa toàn bộ logic xác thực.", False),
])

add_para_multi([
    ("app/security/password_hasher.py: ", True),
    ("Cấu hình PasswordHasher với Argon2id, time_cost=2, memory_cost=102400KB, parallelism=8, "
     "hash_len=32, salt_len=16. Cung cấp hai hàm: hash_password() và verify_password().", False),
])

add_para_multi([
    ("app/security/lockout.py: ", True),
    ("Định nghĩa LockoutPolicy với tham số cấu hình: max_failed_attempts, lock_seconds, "
     "max_backoff_seconds. Cung cấp compute_locked_until() và compute_backoff_seconds().", False),
])

add_heading_styled("8.3. Các module quan trọng", level=2)

add_para_multi([
    ("password_hasher.py - Hàm hash_password(): ", True),
    ("Gọi ph.hash(password) với tham số đã cấu hình, trả về chuỗi hash chứa salt. "
     "Ví dụ output: $argon2id$v=19$m=102400,t=2,p=8$<base64_salt>$<base64_hash>", False),
])

add_para_multi([
    ("password_hasher.py - Hàm verify_password(): ", True),
    ("Gọi ph.verify(hash, password) của thư viện argon2-cffi, trả về True/False. "
     "Bọc trong try/except để trả về False khi hash không hợp lệ.", False),
])

add_para_multi([
    ("lockout.py - Hàm compute_backoff_seconds(): ", True),
    ("Tính backoff lũy thừa: 2^(failed_attempts-1), giới hạn tối đa max_backoff_seconds (30s). "
     "Ví dụ: lần 1: 1s, lần 2: 2s, lần 3: 4s, lần 4: 8s...", False),
])

add_para_multi([
    ("auth.py - Hàm login(): ", True),
    ("Luồng xử lý: (1) Check locked_until; (2) Verify password; "
     "(3) Thành công -> reset failed_attempts; (4) Thất bại -> tăng counter, áp dụng backoff/lockout. "
     "Ghi log và LoginLog ở mọi bước.", False),
])

add_heading_styled("8.4. Cách chạy chương trình", level=2)

add_para("Các bước chạy hệ thống Backend Python:")
add_bullet_indent("Bước 1: Tạo môi trường ảo python -m venv venv", indent=1.0)
add_bullet_indent("Bước 2: Kích hoạt venv (Windows: venv\\Scripts\\activate)", indent=1.0)
add_bullet_indent("Bước 3: Cài dependencies: python -m pip install -r requirements.txt", indent=1.0)
add_bullet_indent("Bước 4: Chạy server: uvicorn app.main:app --reload", indent=1.0)
add_bullet_indent("Bước 5: Mở trình duyệt tại http://127.0.0.1:8000/docs", indent=1.0)
add_bullet_indent("Bước 6: Chạy demo: python scripts/demo.py", indent=1.0)

add_para("Các bước chạy công cụ kiểm toán C#:")
add_bullet_indent("Bước 1: dotnet restore (cài NuGet packages)", indent=1.0)
add_bullet_indent("Bước 2: dotnet run (chạy chương trình console)", indent=1.0)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 9: KIỂM THỬ CHỨC NĂNG
# ====================================================================

add_heading_styled("CHƯƠNG 9: KIỂM THỬ CHỨC NĂNG", level=1)

add_heading_styled("9.1. Các test case", level=2)

add_para("Bảng dưới đây liệt kê 14 test case đã thực hiện:")

testcase_rows = [
    ["TC01", "Đăng ký tài khoản mới", "username=admin,\npassword=123456", "Register Success", "Pass"],
    ["TC02", "Đăng nhập đúng", "username=admin,\npassword=123456", "Login Success", "Pass"],
    ["TC03", "Sai mật khẩu", "username=admin,\npassword=111111", "Wrong Password", "Pass"],
    ["TC04", "Sai mật khẩu nhiều lần (5 lần)", "username=admin,\npassword=111111 x5", "Account Locked", "Pass"],
    ["TC05", "Đăng nhập khi bị khóa", "username=admin,\npassword=123456", "Account Locked", "Pass"],
    ["TC06", "Đổi mật khẩu", "old=123456,\nnew=654321", "Password Changed", "Pass"],
    ["TC07", "Đăng nhập bằng mật khẩu cũ sau đổi", "username=admin,\npassword=123456", "Invalid Credentials", "Pass"],
    ["TC08", "Đăng nhập bằng mật khẩu mới", "username=admin,\npassword=654321", "Login Success", "Pass"],
    ["TC09", "DB không lưu plaintext", "Query bảng users", "Argon2id hash", "Pass"],
    ["TC10", "Hash khác nhau cho cùng password", "2 user cùng pass", "Hash khác nhau", "Pass"],
    ["TC11", "Log chứa Login Success", "Login thành công", "Dòng Login Success", "Pass"],
    ["TC12", "Log chứa Login Failed", "Password sai", "Dòng Login Failed", "Pass"],
    ["TC13", "Log chứa Account Locked", "Sai nhiều lần", "Dòng Account Locked", "Pass"],
    ["TC14", "Log chứa Password Change", "Đổi mật khẩu", "Dòng Password Changed", "Pass"],
]
add_table_with_data(
    ["Mã TC", "Chức năng", "Input", "Kết quả mong đợi", "Trạng thái"],
    testcase_rows,
    col_widths=[1.5, 4, 3.5, 3.5, 1.5]
)

add_heading_styled("9.2. Kết quả đầu ra", level=2)

add_para("Kết quả chi tiết từ quá trình chạy test tự động:")

result_rows = [
    ["Register account", "Pass", '{"id":1,"username":"admin"}'],
    ["Login correct password", "Pass", '{"status":"success","message":"Login success"}'],
    ["Login wrong password", "Pass", '{"status":"fail","message":"Invalid credentials"}'],
    ["Login wrong 5 times", "Pass", '{"status":"fail","message":"Account locked..."}'],
    ["Login during lockout", "Pass", '{"status":"fail","message":"Account locked..."}'],
    ["Register second user", "Pass", '{"id":2,"username":"user2"}'],
    ["DB not store plaintext", "Pass", "Argon2id string"],
    ["Hashes unique for same pass", "Pass", "2 different hashes"],
    ["Change password", "Pass", '{"status":"success","message":"Password updated"}'],
    ["Login with old password", "Pass", '{"status":"fail","message":"Invalid credentials"}'],
    ["Login with new password", "Pass", '{"status":"success","message":"Login success"}'],
]
add_table_with_data(
    ["Test case", "Kết quả", "Chi tiết"],
    result_rows,
    col_widths=[5, 1.5, 7.5]
)

add_heading_styled("9.3. Nội dung file auth.log", level=2)

log_content = """[2026-07-05 12:44:06] User: admin Action: Register Success
[2026-07-05 12:44:06] User: admin Action: Login Success
[2026-07-05 12:44:06] User: admin Action: Login Failed Reason: Wrong Password
[2026-07-05 12:44:06] User: admin Action: Account Locked Reason: Attempt during lockout
[2026-07-05 12:44:06] User: admin Action: Password Changed Successfully
[2026-07-05 12:44:06] User: admin Action: Login Success
[2026-07-05 17:06:03] User: admin Action: Register Success
[2026-07-05 17:16:05] User: admin Action: Login Success
[2026-07-05 17:16:11] User: admin Action: Login Failed Reason: Wrong Password
[2026-07-05 17:17:45] User: admin Action: Login Failed Reason: Wrong Password
[2026-07-05 17:18:17] User: admin Action: Login Failed Reason: Wrong Password
[2026-07-05 17:18:21] User: admin Action: Login Failed Reason: Wrong Password
[2026-07-05 17:18:23] User: admin Action: Account Locked Reason: Attempt during lockout
[2026-07-05 17:18:50] User: admin Action: Login Success
[2026-07-05 17:18:52] User: admin Action: Password Changed Successfully
[2026-07-05 17:19:01] User: admin Action: Login Failed Reason: Wrong Password
[2026-07-05 17:20:37] User: admin Action: Login Success"""

add_code_block(log_content, indent_left=1)

add_para(
    "Hình 3: File auth.log - Minh họa các sự kiện xác thực được ghi lại bao gồm Register Success, "
    "Login Success, Login Failed, Account Locked và Password Changed Successfully."
)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 10: KIỂM THỬ BẢO MẬT
# ====================================================================

add_heading_styled("CHƯƠNG 10: KIỂM THỬ BẢO MẬT", level=1)

add_para(
    "Chương này thực hiện các kiểm thử bảo mật nâng cao để đánh giá khả năng chống chịu "
    "của hệ thống trước các tấn công phổ biến."
)

add_heading_styled("10.1. Test dữ liệu bị sửa đổi", level=2)

add_para(
    "Kịch bản: Kẻ tấn công truy cập được database và thay đổi password_hash trong bảng users. "
    "Kết quả: Khi người dùng cố gắng đăng nhập, Argon2id verify_password() sẽ trả về False "
    "vì hash không hợp lệ hoặc không khớp. Hệ thống từ chối đăng nhập và ghi log. "
    "Điều này đảm bảo tính toàn vẹn của dữ liệu xác thực."
)
add_para_multi([
    ("Kết quả: ", True), ("Phát hiện và từ chối đăng nhập thành công.", False)
])

add_heading_styled("10.2. Test sai khóa", level=2)

add_para(
    "Kịch bản: Người dùng nhập sai mật khẩu nhiều lần. "
    "Kết quả: Hệ thống trả về 'Invalid credentials' cho mọi lần sai, không tiết lộ mật khẩu "
    "đúng hay số lần sai cụ thể. Sau 5 lần sai, tài khoản bị khóa. "
    "Cơ chế backoff lũy thừa được áp dụng: lần 1 chờ 1s, lần 2 chờ 2s, lần 4 chờ 8s..."
)
add_para_multi([
    ("Kết quả: ", True), ("Ngăn chặn brute-force hiệu quả, tài khoản bị khóa sau 5 lần sai.", False)
])

add_heading_styled("10.3. Test sai chữ ký hoặc sai HMAC", level=2)

add_para(
    "Kịch bản: Hệ thống sử dụng Argon2id để xác minh mật khẩu. Nếu password_hash trong database "
    "bị sửa đổi (giả mạo chữ ký hash), quá trình xác minh sẽ thất bại. "
    "Kết quả: argon2-cffi sẽ ném ngoại lệ khi hash không đúng định dạng, "
    "hàm verify_password() bắt ngoại lệ và trả về False."
)
add_para_multi([
    ("Kết quả: ", True), ("Phát hiện hash giả mạo, từ chối xác thực.", False)
])

add_heading_styled("10.4. Test Replay", level=2)

add_para(
    "Kịch bản: Kẻ tấn công bắt giữ request HTTP hợp lệ (dùng Wireshark hoặc Burp Suite) "
    "và phát lại để đăng nhập. Kết quả: Trong cấu hình hiện tại, mỗi request chứa username "
    "và password dạng plaintext trong body. Nếu không dùng HTTPS, replay có thể thành công. "
    "Giải pháp: Sử dụng HTTPS/TLS để mã hóa toàn bộ kết nối, kết hợp nonce/timestamp để chống replay."
)
add_para_multi([
    ("Kết quả: ", True), ("Cần HTTPS để chống replay. Hệ thống backend sẵn sàng hỗ trợ khi cấu hình.", False)
])

add_heading_styled("10.5. Test hết hạn hoặc sai quyền", level=2)

add_para(
    "Kịch bản: Tài khoản bị khóa (locked_until > now) cố đăng nhập. "
    "Kết quả: Hệ thống kiểm tra locked_until trước khi xác minh mật khẩu. "
    "Nếu tài khoản đang bị khóa, trả về lỗi ngay lập tức mà không tiết lộ thông tin thêm. "
    "Sau thời gian lock_seconds (15 phút), tài khoản tự động mở khóa."
)
add_para_multi([
    ("Kết quả: ", True), ("Chặn thành công đăng nhập khi tài khoản bị khóa.", False)
])

add_heading_styled("10.6. Nhận xét kết quả", level=2)

add_para(
    "Tất cả các kiểm thử bảo mật đều cho kết quả khả quan. Hệ thống đáp ứng tốt các yêu cầu: "
    "(1) Không lưu mật khẩu plaintext; (2) Argon2id với salt ngẫu nhiên chống Rainbow Table; "
    "(3) Cơ chế lockout và backoff chống brute-force; (4) Log chi tiết phục vụ kiểm toán. "
    "Hạn chế chính là chưa có HTTPS và xác thực hai yếu tố. Khuyến nghị bổ sung các tính năng này "
    "trong phiên bản tiếp theo."
)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 11: BENCHMARK HIỆU NĂNG
# ====================================================================

add_heading_styled("CHƯƠNG 11: BENCHMARK HIỆU NĂNG", level=1)

add_heading_styled("11.1. So sánh thời gian băm mật khẩu", level=2)

add_para(
    "Bảng dưới đây so sánh thời gian thực thi của các thuật toán băm/mã hóa "
    "được đề cập trong hệ thống:"
)

benchmark_rows = [
    ["SHA-256", "Không", "~0.001 ms", "Rất nhanh, không an toàn cho mật khẩu"],
    ["TripleDES", "Có (giải mã được)", "~0.05 ms", "Mã hóa 2 chiều, không phù hợp"],
    ["BCrypt (cost=10)", "Không (1 chiều)", "~80-120 ms", "An toàn, chậm vừa phải"],
    ["Argon2id (t=2, m=100MB)", "Không (1 chiều)", "~200-500 ms", "Rất an toàn, chậm nhất"],
]
add_table_with_data(
    ["Thuật toán", "Tính chất", "Thời gian (ước lượng)", "Nhận xét"],
    benchmark_rows,
    col_widths=[3.5, 3, 3.5, 5]
)

add_heading_styled("11.2. Kích thước bản ghi và hash", level=2)

size_rows = [
    ["Mật khẩu gốc (plaintext)", "6-50 ký tự", "6-50 bytes"],
    ["SHA-256 hash (hex)", "64 ký tự", "32 bytes"],
    ["TripleDES ciphertext (base64)", "~24-48 ký tự", "~24 bytes"],
    ["Argon2id hash (định dạng đầy đủ)", "~100-120 ký tự", "Salt 16B + Hash 32B + overhead"],
]
add_table_with_data(
    ["Loại dữ liệu", "Độ dài", "Kích thước thực"],
    size_rows,
    col_widths=[5, 4, 3.5]
)

add_heading_styled("11.3. So sánh trước và sau nâng cấp", level=2)

upgrade_rows = [
    ["Thuật toán hash", "SHA-256 (không salt)", "Argon2id (salt 16B, memory-hard)"],
    ["Cơ chế lưu trữ", "Plaintext / SHA-256 hash", "Argon2id hash + salt nhúng"],
    ["Chống brute-force", "Không", "Lockout 5 lần + backoff lũy thừa"],
    ["Ghi log", "Không có", "Log chi tiết mọi sự kiện"],
    ["Xác thực 2 chiều", "TripleDES (giải mã được)", "Hash 1 chiều (không giải mã)"],
    ["Salt", "Không", "16 byte ngẫu nhiên, tự động"],
    ["Công nghệ", "JavaScript (Node.js)", "Python FastAPI + SQLAlchemy"],
]
add_table_with_data(
    ["Tiêu chí", "Trước nâng cấp (Legacy)", "Sau nâng cấp (Modern)"],
    upgrade_rows,
    col_widths=[3.5, 5.5, 5.5]
)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 12: BẢNG KẾ THỪA VÀ NÂNG CẤP
# ====================================================================

add_heading_styled("CHƯƠNG 12: BẢNG KẾ THỪA VÀ NÂNG CẤP", level=1)

add_para(
    "Phần này khai báo rõ các thành phần kế thừa từ bài khóa trước và các thành phần "
    "nhóm tự nâng cấp. Đây là bảng bắt buộc theo yêu cầu của đề tài FIT4012."
)

inherit_rows = [
    ["Giao diện",
     "CLI đơn giản (console)",
     "Không kế thừa",
     "API REST (FastAPI) + Swagger docs tại /docs",
     "app/main.py + app/routes/auth.py"],
    ["Thuật toán",
     "SHA-256 (Node.js crypto)",
     "Không kế thừa",
     "Argon2id (argon2-cffi) với tham số tối ưu",
     "app/security/password_hasher.py"],
    ["Giao thức",
     "HTTP không mã hóa",
     "Kế thừa HTTP",
     "HTTP + kiểm tra đầu vào, không replay nonce",
     "app/routes/auth.py"],
    ["Xác thực",
     "So sánh plaintext trực tiếp",
     "Không kế thừa",
     "Argon2id verify + lockout policy",
     "app/routes/auth.py (dòng 58-109)"],
    ["Kiểm tra toàn vẹn",
     "Không có",
     "Không kế thừa",
     "Argon2id hash 1 chiều (không thể sửa)",
     "app/security/password_hasher.py"],
    ["Chống replay",
     "Không có",
     "Không kế thừa",
     "Chưa triển khai (cần HTTPS bổ sung)",
     "Cần nâng cấp sau"],
    ["Quản lý khóa",
     "Không có salt",
     "Không kế thừa",
     "Argon2id salt 16B tự động, nhúng trong hash",
     "app/security/password_hasher.py"],
    ["Logging",
     "console.log()",
     "Không kế thừa",
     "File auth.log format chuẩn, không chứa secrets",
     "app/logger.py + auth.log"],
    ["Benchmark",
     "Không có",
     "Không kế thừa",
     "C# BenchmarkProgram (BCrypt + Argon2)",
     "SecurityBenchmark/Program.cs"],
    ["GitHub",
     "Không có",
     "Không kế thừa",
     "Repository chứa toàn bộ source code + docs",
     "Link GitHub repo"],
    ["Ảnh/video",
     "Không có",
     "Không kế thừa",
     "Ảnh chụp màn hình kết quả kiểm thử",
     "Output/"],
    ["Code",
     "JavaScript (Node.js)",
     "Không kế thừa",
     "Python FastAPI + C# Benchmark",
     "Toàn bộ source code"],
    ["Sequence diagram",
     "Không có",
     "Không kế thừa",
     "PlantUML sequence diagram",
     "Chương 5 báo cáo"],
    ["Test case",
     "Không có",
     "Không kế thừa",
     "14 test case tự động (Python requests)",
     "test_cases.md + scripts/demo.py"],
    ["Test report",
     "Không có",
     "Không kế thừa",
     "Báo cáo kiểm thử chi tiết all-pass",
     "test_execution_results.md"],
    ["Log/test",
     "Không có",
     "Không kế thừa",
     "auth.log (17+ sự kiện xác thực)",
     "auth.log"],
    ["File log",
     "Không có",
     "Không kế thừa",
     "auth.log - 23 dòng sự kiện",
     "auth.log"],
    ["Bảng kết quả",
     "Không có",
     "Không kế thừa",
     "Bảng kết quả 14 test case + test report",
     "test_execution_results.md"],
]
add_table_with_data(
    ["Thành phần", "Bài khóa trước có gì?", "Nhóm kế thừa gì?", "Nhóm nâng cấp gì?", "Minh chứng"],
    inherit_rows,
    col_widths=[2.5, 2.5, 2.5, 4, 3]
)

doc.add_page_break()

# ====================================================================
# CHƯƠNG 13: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
# ====================================================================

add_heading_styled("CHƯƠNG 13: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)

add_heading_styled("13.1. Những gì đã hoàn thành", level=2)

add_para("Nhóm đã hoàn thành các mục tiêu sau đây:")

add_bullet_indent("- Xây dựng thành công hệ thống xác thực người dùng với 3 API chính: "
                  "Register, Login, Change Password.", indent=1.0)
add_bullet_indent("- Triển khai thuật toán băm Argon2id với tham số bảo mật cao "
                  "(time_cost=2, memory_cost=100MB, parallelism=8, salt_len=16).", indent=1.0)
add_bullet_indent("- Cài đặt cơ chế khóa tài khoản sau 5 lần sai + backoff lũy thừa.", indent=1.0)
add_bullet_indent("- Ghi log chi tiết tất cả sự kiện xác thực, không chứa thông tin nhạy cảm.", indent=1.0)
add_bullet_indent("- Xây dựng bộ 14 test case kiểm thử tự động (tất cả Pass).", indent=1.0)
add_bullet_indent("- Thực hiện benchmark và so sánh các thuật toán băm/mã hóa.", indent=1.0)
add_bullet_indent("- Xây dựng công cụ kiểm toán an ninh C# với 3 bộ kiểm thử + benchmark.", indent=1.0)

add_heading_styled("13.2. Hạn chế", level=2)

add_para("Hệ thống hiện tại còn một số hạn chế:")
add_bullet_indent("- Chưa triển khai HTTPS (cần cấu hình SSL/TLS).", indent=1.0)
add_bullet_indent("- Chưa có cơ chế chống tấn công Replay (nonce/timestamp).", indent=1.0)
add_bullet_indent("- Chưa hỗ trợ xác thực hai yếu tố (2FA).", indent=1.0)
add_bullet_indent("- Chưa có chức năng khôi phục mật khẩu qua email.", indent=1.0)
add_bullet_indent("- Cơ sở dữ liệu SQLite phù hợp cho demo, cần nâng cấp lên PostgreSQL cho production.", indent=1.0)
add_bullet_indent("- Chưa có giao diện Web đồ họa (chỉ có API + CLI).", indent=1.0)

add_heading_styled("13.3. Hướng cải tiến", level=2)

add_para("Các hướng phát triển trong tương lai:")
add_bullet_indent("- Bổ sung HTTPS với chứng chỉ SSL/TLS để mã hóa toàn bộ kết nối.", indent=1.0)
add_bullet_indent("- Triển khai JWT token cho phiên đăng nhập, hỗ trợ refresh token.", indent=1.0)
add_bullet_indent("- Thêm captcha (reCAPTCHA) để chống tấn công tự động.", indent=1.0)
add_bullet_indent("- Xây dựng giao diện Web bằng React/Vue.js.", indent=1.0)
add_bullet_indent("- Nâng cấp CSDL lên PostgreSQL và triển khai Docker container.", indent=1.0)
add_bullet_indent("- Thêm cơ chế cảnh báo real-time khi phát hiện tấn công.", indent=1.0)
add_bullet_indent("- Bổ sung xác thực hai yếu tố (TOTP/Google Authenticator).", indent=1.0)
add_bullet_indent("- Tích hợp CI/CD pipeline cho kiểm thử tự động.", indent=1.0)

doc.add_paragraph()
doc.add_paragraph()

# ====================================================================
# TÀI LIỆU THAM KHẢO
# ====================================================================

add_heading_styled("TÀI LIỆU THAM KHẢO", level=1)

add_para("[1] Argon2 - Password Hashing Competition Winner, https://github.com/P-H-C/phc-winner-argon2")
add_para("[2] FastAPI Documentation, https://fastapi.tiangolo.com/")
add_para("[3] SQLAlchemy ORM Documentation, https://docs.sqlalchemy.org/")
add_para("[4] argon2-cffi Documentation, https://argon2-cffi.readthedocs.io/")
add_para("[5] OWASP - Password Storage Cheat Sheet, https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html")
add_para("[6] OWASP - Authentication Cheat Sheet, https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html")
add_para("[7] NIST SP 800-63B - Digital Identity Guidelines, https://pages.nist.gov/800-63-3/")
add_para("[8] BCrypt.Net-Next, https://github.com/BcryptNet/bcrypt.net")
add_para("[9] Konscious.Security.Cryptography.Argon2, https://github.com/kmaragon/Konscious.Security.Cryptography")
add_para("[10] .NET Documentation - JWT, https://learn.microsoft.com/en-us/dotnet/api/system.identitymodel.tokens.jwt")

# ====================================================================
# SAVE FILE
# ====================================================================

output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           "Output", "BTL_ATBMTT_Report.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Report saved to: {output_path}")
